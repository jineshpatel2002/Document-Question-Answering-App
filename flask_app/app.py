from flask import Flask, request, render_template, redirect, url_for
from werkzeug.utils import secure_filename
import os
from langchain_community.document_loaders import PyPDFLoader
from docx import Document
from langchain_chroma import Chroma
from langchain_community.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain_text_splitters import CharacterTextSplitter
from langchain_core.documents import Document as customDocument
from langchain_core.prompts import ChatPromptTemplate
from transformers import pipeline

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Function to read and split PDF
def read_pdf(file_path):
    loader = PyPDFLoader(file_path)
    documents = loader.load_and_split()
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = text_splitter.split_documents(documents)
    return docs

# Function to read DOCX
def read_doc(file_path):
    doc = Document(file_path)
    text = "".join([paragraph.text + "\n" for paragraph in doc.paragraphs])
    document = customDocument(page_content=text, metadata={"source": file_path})
    return [document]

# Function to generate embeddings and store in ChromaDB
def generate_embeddings_and_store_chromadb(docs):
    embedding_function = SentenceTransformerEmbeddings(model_name="all-MiniLM-L6-v2")
    db = Chroma.from_documents(docs, embedding_function)
    return db

# Custom chain class
class CustomChain:
    def __init__(self, retriever, prompt, llm):
        self.retriever = retriever
        self.prompt = prompt
        self.llm = llm

    def invoke(self, question):
        context_docs = self.retriever.invoke(question)
        context = "\n\n".join([doc.page_content for doc in context_docs]) if context_docs else ""
        input_text = f"Question: {question}\n\nContext: {context}\n\n Answer in detail:"
        response = self.llm(input_text, max_length=500, num_return_sequences=1, do_sample=False)[0]['generated_text']
        return response

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    if file:
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        if file_path.endswith('.pdf'):
            docs = read_pdf(file_path)
        elif file_path.endswith('.docx'):
            docs = read_doc(file_path)
        else:
            return "Unsupported file format", 400
        
        vector_store = generate_embeddings_and_store_chromadb(docs)
        retriever = vector_store.as_retriever(search_type="similarity", search_kwargs={"k": 3})
        llm = pipeline("text2text-generation", model="google/flan-t5-large")
        
        prompt_message = """
        Please provide a detailed answer to the following question using the provided context.
        
        {question}
        
        Context:
        {context}
        """
        
        prompt = ChatPromptTemplate.from_messages([("human", prompt_message)])
        global rag_chain
        rag_chain = CustomChain(retriever=retriever, prompt=prompt, llm=llm)
        
        return redirect(url_for('ask_question'))

@app.route('/ask', methods=['GET', 'POST'])
def ask_question():
    if request.method == 'POST':
        question = request.form['question']
        response = rag_chain.invoke(question)
        return render_template('result.html', question=question, answer=response)
    return render_template('ask.html')

if __name__ == '__main__':
    app.run()
